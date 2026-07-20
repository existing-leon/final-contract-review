"""Word 文档解析（python-docx）。"""
from app.core.constants import ParseStatus
from app.core.exceptions import BizException, Errors


def parse_docx(file_path: str) -> dict:
    try:
        import docx
    except ImportError as e:
        raise BizException(*Errors.OCR_FAILED, data={"reason": "未安装 python-docx"}) from e

    try:
        doc = docx.Document(file_path)
    except Exception as e:
        raise BizException(*Errors.OCR_FAILED, data={"reason": str(e)}) from e

    parts = []
    # 正文段落
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts).strip()
    if not full_text:
        raise BizException(*Errors.DOC_EMPTY, data={"reason": "文档内容为空"})

    return {
        "text": full_text,
        "pages": [{"page": 1, "text": full_text}],
        "parse_status": ParseStatus.SUCCESS,
        "parse_error": None,
    }
